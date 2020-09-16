import logging
import json
import filetype
import base64
import io
import re

from docx import Document
from pdfreader import SimplePDFViewer

import azure.functions as func

# Parse Attachment
# 
# Takes a file (base64 encoded binary) and validates it as a TSA RTS Risk Assessment
#  - Supports DOCX and PDF formats
#  - Attempts to validate document before processing
#  - Trys to fail safe!

def main(req: func.HttpRequest) -> func.HttpResponse:

    logging.info('Processing a file to extract details')

    bodyObj = json.loads(req.get_body())
    
    fileContents = bodyObj["fileContents"]

    docObj = {}

    #No file content
    if not fileContents:
       docObj["status"] = "error"
       docObj["error"] = "File contents not provided"
       return func.HttpResponse(json.dumps(docObj))

    #Process/parse file
    if fileContents:

        #Base64 Decode the File
        decodedFile = base64.standard_b64decode(fileContents)

        downloadedfile = io.BytesIO(decodedFile)

        kind = filetype.guess(decodedFile)

        #If Zip (DOCX Hack, eurgh. CBA to write custom parser, don't judge me.)
        if kind.mime == "application/zip":

            try:
                document: Document = Document(downloadedfile)
                headingText = document.sections[0].header.paragraphs[0].text

                if headingText == "Covid-19 restarting face to face Scouting risk assessment":
                    logging.info("Found DOCX Risk Assessment!")

                    docObj['section'] = document.tables[0].cell(0,1).text
                    docObj['date'] = document.tables[0].cell(0,3).text
                    docObj['author'] = document.tables[0].cell(0,5).text
                    docObj['level'] = document.tables[0].cell(0,7).text

                    docObj["status"] = "ok"
                    docObj["risk-assessment"] = True

                else:
                    #File Header not matched
                    docObj["status"] = "ok"
                    docObj["risk-assessment"] = False

            except Exception as e:
                docObj["status"] = "error"
                docObj["risk-assessment"] = False
                docObj["error"] = e.__cause__

        #If PDF
        elif kind.mime == "application/pdf":

            try:
                viewer = SimplePDFViewer(downloadedfile)
                viewer.render()
                
                textContent = "".join(viewer.canvas.strings)

                if textContent.startswith("Covid-19 restarting face to face Scouting risk assessment"):
                    logging.info("Found PDF Risk Assessment!")

                    # Extraction via hacky regex's. Docx is much better.
                    #Section
                    sectionMatch = re.search("Name of Section or Activity(.*)Date of", textContent)
                    if sectionMatch:
                        docObj['section'] = sectionMatch.group(1).strip()
                    #Date
                    dateMatch = re.search("Date of  risk assessment(.*)Name", textContent)
                    if dateMatch:
                        docObj['date'] = dateMatch.group(1).strip()
                    #Author
                    authorMatch = re.search("Name of who undertook this risk assessment(.*)COVID-19", textContent)
                    if authorMatch:
                        docObj['author'] = authorMatch.group(1).strip()
                    #Level
                    levelMatch = re.search("COVID-19 readiness level transition(.*)Hazard Identified", textContent)
                    if levelMatch:
                        docObj['level'] = levelMatch.group(1).strip()

                    docObj["status"] = "ok"
                    docObj["risk-assessment"] = True

                else:
                    #File Header not matched
                    docObj["status"] = "ok"
                    docObj["risk-assessment"] = False

            except Exception as e:
                docObj["status"] = "error"
                docObj["risk-assessment"] = False
                docObj["error"] = e.__cause__
       
        #Neither, unprocessable
        else:
            docObj["status"] = "ok"
            docObj["risk-assessment"] = False

        #Return
        headers = {"Content-Type": "application/json"}

        return func.HttpResponse(json.dumps(docObj), headers=headers)

    else:
        return func.HttpResponse(
             "No Body Processed",
             status_code=500
        )
